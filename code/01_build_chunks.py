import json  # 导入 json，用来把 chunk 写成 jsonl 文件。
import hashlib  # 导入 hashlib，用来为源 Word 路径生成稳定的文件标识。
import re  # 导入 re，用正则识别章节标题和清理空白字符。
from collections import Counter  # 导入 Counter，用来统计章节、文件、chunk 数量。
from pathlib import Path  # 导入 Path，用来处理文件路径。

from docx import Document as WordDocument  # 从 python-docx 导入 WordDocument，用来读取 docx 正文。
from langchain_core.documents import Document  # 从 LangChain 导入 Document，这是 LangChain 统一承载文本和元数据的对象。
from langchain_text_splitters import RecursiveCharacterTextSplitter  # 从 LangChain 导入递归切分器，让长文本按段落边界优先切开。

from config import CHUNK_OVERLAP  # 从配置文件读取 chunk 重叠长度。
from config import CHUNK_SEPARATORS  # 从配置文件读取固定切分边界，保证建库和版本指纹使用同一套规则。
from config import CHUNK_SIZE  # 从配置文件读取 chunk 目标长度。
from config import CHUNKS_JSONL  # 从配置文件读取 chunk 输出路径。
from config import COLLECTION_NAME  # 从配置文件读取知识库名称。
from config import INDEX_ROOT  # 从配置文件读取索引目录。
from config import MANIFEST_JSON  # 从配置文件读取清单路径。
from config import SOURCE_ROOT  # 从配置文件读取源文档目录。
from index_health import artifact_info  # 导入 artifact 摘要函数，让 chunk 构建阶段也能留下可验证的文件指纹。
from index_health import create_manifest  # 导入完整索引清单创建函数。
from index_health import write_manifest  # 导入统一清单写入函数。


def normalize_text(text: str) -> str:  # 定义文本清洗函数，输入原始文本，输出更适合检索的文本。
    text = text.replace("\u3000", " ")  # 把中文全角空格替换成普通空格，避免检索时出现奇怪空白。
    text = re.sub(r"[ \t]+", " ", text)  # 把连续空格和制表符合并成一个空格。
    text = re.sub(r"\n{3,}", "\n\n", text)  # 把三行以上空行压缩成两行，保留段落感但不浪费上下文。
    return text.strip()  # 去掉开头和结尾的空白，返回干净文本。


def read_docx_text(path: Path) -> str:  # 定义读取 Word 文件的函数，输入 docx 路径，输出纯文本。
    word = WordDocument(str(path))  # 打开 Word 文件，显式传字符串路径兼容 python-docx 类型声明。
    blocks: list[str] = []  # 准备一个列表，用来按顺序保存段落和表格文字。
    for paragraph in word.paragraphs:  # 遍历 Word 里的每一个普通段落。
        text = normalize_text(paragraph.text)  # 清洗当前段落文字。
        if text:  # 如果段落不是空的，才进入结果。
            blocks.append(text)  # 保存当前段落。
    for table in word.tables:  # 遍历 Word 里的每一个表格。
        for row in table.rows:  # 遍历表格中的每一行。
            cells = [normalize_text(cell.text) for cell in row.cells]  # 读取并清洗这一行的每个单元格。
            cells = [cell for cell in cells if cell]  # 去掉空单元格，避免产生无意义分隔符。
            if cells:  # 如果这一行表格有文字，才保存。
                blocks.append(" | ".join(cells))  # 用竖线连接单元格，让表格在纯文本里仍然有结构。
    return normalize_text("\n\n".join(blocks))  # 用空行连接所有块，并做最后一次清洗。


def infer_chapter_parts(path: Path) -> dict:  # 定义从文件路径推断章节元数据的函数。
    relative = path.relative_to(SOURCE_ROOT)  # 计算文件相对源文档目录的路径。
    parts = list(relative.parts)  # 把相对路径拆成目录名和文件名列表。
    filename = path.stem  # 取文件名但不要 .docx 后缀，通常它就是章节名。
    chapter = parts[0] if len(parts) >= 2 else filename  # 如果文件在章节文件夹内，就把第一层目录当章名。
    section = filename  # 默认把当前文件名当小节名。
    match = re.match(r"^(\d+(?:\.\d+)+)\s*(.*)$", filename)  # 尝试识别类似 5.2.1 标题的编号。
    section_number = match.group(1) if match else ""  # 如果识别成功，就保存小节编号。
    section_title = match.group(2).strip() if match else filename  # 如果识别成功，就保存去掉编号后的标题。
    return {  # 返回一个元数据字典。
        "relative_path": str(relative),  # 保存相对路径，方便引用时展示。
        "source_file_id": hashlib.sha1(str(relative).encode("utf-8"), usedforsecurity=False).hexdigest()[:12],  # 这里只把 SHA-1 当稳定 ID，不用于安全校验，显式关闭安全用途告警。
        "chapter": chapter,  # 保存章名。
        "section": section,  # 保存小节名。
        "section_number": section_number,  # 保存小节编号。
        "section_title": section_title,  # 保存小节标题。
    }  # 元数据字典到这里结束。


def build_contextual_text(metadata: dict, text: str) -> str:  # 定义给正文加上下文前缀的函数。
    header = "\n".join(  # 组装一个短标题区，让 embedding 能看到章节背景。
        [  # 标题区每一行都是检索上下文。
            f"书名：{COLLECTION_NAME}",  # 放入书名，帮助模型知道资料来源。
            f"章节：{metadata['chapter']}",  # 放入章名，帮助检索命中章节主题。
            f"小节：{metadata['section']}",  # 放入小节名，帮助检索命中具体知识点。
            "正文：",  # 标明下面开始是原文正文。
        ]  # 标题区列表结束。
    )  # 标题区字符串组装完成。
    return f"{header}\n{text}"  # 返回带上下文前缀的完整文本。


def load_documents() -> list[Document]:  # 定义加载所有 Word 文件并转换成 LangChain Document 的函数。
    docs: list[Document] = []  # 准备一个列表，保存所有 LangChain Document。
    for path in sorted(SOURCE_ROOT.rglob("*.docx")):  # 递归扫描源目录下所有 docx 文件。
        text = read_docx_text(path)  # 读取当前 Word 文件文本。
        if not text:  # 如果这个 Word 没有有效文本，就跳过。
            continue  # 跳到下一个文件。
        metadata = infer_chapter_parts(path)  # 从路径和文件名推断章节元数据。
        metadata["source_docx"] = str(path)  # 保存原始 Word 文件绝对路径，方便追溯。
        metadata["collection"] = COLLECTION_NAME  # 保存知识库名称，方便未来多知识库过滤。
        contextual_text = build_contextual_text(metadata, text)  # 给正文加上书名、章节、小节上下文。
        docs.append(Document(page_content=contextual_text, metadata=metadata))  # 创建 LangChain Document 并放入列表。
    return docs  # 返回所有文档对象。


def split_documents(docs: list[Document]) -> list[Document]:  # 定义把长文档切成 chunk 的函数。
    splitter = RecursiveCharacterTextSplitter(  # 创建 LangChain 递归文本切分器。
        chunk_size=CHUNK_SIZE,  # 设置每块目标长度。
        chunk_overlap=CHUNK_OVERLAP,  # 设置相邻块重叠长度。
        separators=list(CHUNK_SEPARATORS),  # 设置优先按段落、句号、分号等自然边界切。
    )  # 切分器创建完成。
    chunks = splitter.split_documents(docs)  # 使用 LangChain 切分器把所有文档切成 chunk。
    source_texts = {doc.metadata["source_file_id"]: doc.page_content.split("\n正文：\n", 1)[-1] for doc in docs}  # 按源文件 ID 保存去掉检索前缀的正文，用来计算源 Word 正文中的字符范围。
    grouped_chunks: dict[str, list[Document]] = {}  # 准备按源文件分组，计算小节内顺序和总段数。
    for chunk in chunks:  # 先遍历切分结果建立源文件分组。
        grouped_chunks.setdefault(chunk.metadata["source_file_id"], []).append(chunk)  # 把当前 chunk 放入所属 Word 文件。
    for index, chunk in enumerate(chunks):  # 遍历每一个 chunk，给它补充唯一编号和稳定定位信息。
        chunk.metadata["chunk_index"] = index  # 保存全局 chunk 序号。
        chunk.metadata["chunk_id"] = f"xg-{index:06d}"  # 保存稳定 chunk_id，后续向量库和引用都用它。
        chunk.metadata["text_length"] = len(chunk.page_content)  # 保存 chunk 长度，方便调参分析。
        file_chunks = grouped_chunks[chunk.metadata["source_file_id"]]  # 取出当前源 Word 的全部 chunk。
        section_index = file_chunks.index(chunk)  # 计算当前 chunk 在源 Word 小节中的顺序。
        chunk.metadata["section_chunk_index"] = section_index  # 保存从 0 开始的小节 chunk 序号。
        chunk.metadata["section_chunk_count"] = len(file_chunks)  # 保存该小节总 chunk 数，展示时转换成从 1 开始的序号。
        body = chunk.page_content.split("\n正文：\n", 1)[-1]  # 去掉首个 chunk 的书名、章节和小节前缀，只保留可定位正文。
        source_text = source_texts[chunk.metadata["source_file_id"]]  # 取出当前 Word 的完整正文。
        anchor = body[:120].strip()  # 取正文开头作为稳定锚点，前后文检索失败时仍可人工搜索。
        search_start = max(0, file_chunks[section_index - 1].metadata.get("source_char_start", 0) - CHUNK_OVERLAP) if section_index else 0  # 从上一个 chunk 附近开始搜索，兼容重叠切分。
        char_start = source_text.find(anchor, search_start) if anchor else -1  # 在源正文中定位当前 chunk 的开头。
        if char_start < 0 and anchor:  # 如果重叠或切分清洗导致局部搜索失败。
            char_start = source_text.find(anchor)  # 退回全局搜索，保证定位字段尽量可用。
        char_end = char_start + len(body) if char_start >= 0 else -1  # 根据正文长度计算结束位置，-1 表示无法自动定位。
        chunk.metadata["source_char_start"] = char_start  # 保存源 Word 正文中的起始字符位置。
        chunk.metadata["source_char_end"] = char_end  # 保存源 Word 正文中的结束字符位置。
        chunk.metadata["source_anchor"] = re.sub(r"\s+", " ", anchor)[:80]  # 保存短锚点，方便用户在 Word 中搜索原文。
    return chunks  # 返回切好的 chunk 列表。


def save_chunks(chunks: list[Document]) -> None:  # 定义把 chunk 保存到 jsonl 的函数。
    INDEX_ROOT.mkdir(parents=True, exist_ok=True)  # 确保索引目录存在。
    with CHUNKS_JSONL.open("w", encoding="utf-8") as file:  # 打开 jsonl 文件准备写入。
        for chunk in chunks:  # 遍历每一个 chunk。
            record = {"text": chunk.page_content, "metadata": chunk.metadata}  # 把正文和元数据组织成一条记录。
            file.write(json.dumps(record, ensure_ascii=False) + "\n")  # 写入一行 JSON，ensure_ascii=False 保留中文。


def save_manifest(docs: list[Document], chunks: list[Document]) -> None:  # 定义保存建库统计清单的函数。
    chapter_counter = Counter(doc.metadata["chapter"] for doc in docs)  # 统计每章有多少 Word 文件。
    manifest = create_manifest(len(docs), len(chunks))  # 创建包含源文件、切分配置、embedding 和索引版本的完整清单。
    manifest["chapters"] = dict(chapter_counter)  # 保留每章 Word 文件统计，便于人工查看教材覆盖情况。
    manifest["artifacts"] = {"chunks_jsonl": artifact_info(CHUNKS_JSONL)}  # 记录 chunk 文件的大小和 SHA256，供健康检查核对。
    manifest["build_status"]["chunks"] = CHUNKS_JSONL.exists()  # 标记 chunk 构建阶段是否已经完成。
    write_manifest(manifest)  # 用统一函数写入完整清单。


def main() -> None:  # 定义主函数，让脚本从这里开始执行。
    docs = load_documents()  # 第一步：读取所有 Word，得到 LangChain Document。
    chunks = split_documents(docs)  # 第二步：把长文档切成适合检索的 chunk。
    save_chunks(chunks)  # 第三步：把 chunk 明细保存成 jsonl。
    save_manifest(docs, chunks)  # 第四步：保存本次建库清单。
    print(f"读取 Word 文件：{len(docs)} 个")  # 打印读取到的 Word 数量。
    print(f"生成文本块：{len(chunks)} 个")  # 打印切出的 chunk 数量。
    print(f"chunk 文件：{CHUNKS_JSONL}")  # 打印 chunk 文件路径。
    print(f"清单文件：{MANIFEST_JSON}")  # 打印清单文件路径。


if __name__ == "__main__":  # 判断当前文件是否被直接运行。
    main()  # 如果是直接运行，就执行主函数。
